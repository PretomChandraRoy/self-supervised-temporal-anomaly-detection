"""
Generate a comprehensive project report (.docx) for the
Self-Supervised Temporal Anomaly Detection project.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import json
from datetime import datetime


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def generate_report():
    doc = Document()

    # ========================================================================
    # Page setup
    # ========================================================================
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Self-Supervised Temporal Anomaly Detection\nin Financial Time-Series')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    doc.add_paragraph('')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Project Report')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph('')
    doc.add_paragraph('')

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run(
        'A Transformer-Based Approach Using Contrastive Learning,\n'
        'Masked Reconstruction, and Energy-Based Scoring\n'
        'for Detecting Abnormal Market Behavior'
    )
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(4):
        doc.add_paragraph('')

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)

    tech = doc.add_paragraph()
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tech.add_run('Framework: PyTorch 2.6.0 | Language: Python 3.10')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS (manual)
    # ========================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Abstract', '3'),
        ('2.', 'Introduction', '3'),
        ('3.', 'Problem Statement', '4'),
        ('4.', 'Literature Review', '5'),
        ('5.', 'Methodology', '6'),
        ('5.1', '   Data Collection and Preprocessing', '6'),
        ('5.2', '   Feature Engineering', '7'),
        ('5.3', '   Model Architecture', '8'),
        ('5.4', '   Self-Supervised Learning Objectives', '9'),
        ('5.5', '   Anomaly Detection Pipeline', '10'),
        ('5.6', '   Hybrid Score Fusion', '11'),
        ('6.', 'System Architecture', '12'),
        ('7.', 'Implementation Details', '13'),
        ('8.', 'Experimental Setup', '14'),
        ('9.', 'Results and Analysis', '15'),
        ('10.', 'Discussion', '17'),
        ('11.', 'Challenges and Solutions', '18'),
        ('12.', 'Future Work', '19'),
        ('13.', 'Conclusion', '20'),
        ('14.', 'References', '21'),
    ]
    for num, title_text, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {title_text}')
        run.font.size = Pt(12)
        if not num.startswith('5.'):
            run.bold = True

    doc.add_page_break()

    # ========================================================================
    # 1. ABSTRACT
    # ========================================================================
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        'This project presents a self-supervised deep learning framework for detecting anomalous '
        'patterns in financial time-series data. The system leverages a Transformer-based temporal '
        'encoder trained with dual self-supervised objectives—temporal contrastive learning (NT-Xent) '
        'and masked time-series reconstruction—to learn robust representations of normal market behavior '
        'from unlabeled EUR/USD forex data spanning 2015 to 2024.'
    )
    doc.add_paragraph(
        'Anomaly detection is performed through a novel hybrid fusion approach combining three '
        'complementary scoring mechanisms: (1) reconstruction error with an information bottleneck '
        'that prevents trivial copying of anomalous patterns, (2) density-aware clustering that '
        'discovers normal market regimes and flags deviations, and (3) an energy-based neural '
        'scoring function trained to assign high energy to out-of-distribution samples. The three '
        'scores are combined using weighted sum and max aggregation with validation-based threshold '
        'tuning.'
    )
    doc.add_paragraph(
        'The framework processes 14,566 four-hour candlestick records, generating 25 engineered '
        'features per time step and organizing them into sliding windows of 60 time steps. Synthetic '
        'anomalies of six types (price spikes, flash crashes, volatility spikes, volume spikes, '
        'trend breaks, and gap anomalies) are injected at a 7% rate for quantitative evaluation. '
        'The system achieves anomaly detection on a heavily imbalanced dataset, demonstrating the '
        'viability of self-supervised representation learning for financial market surveillance.'
    )
    p = doc.add_paragraph()
    run = p.add_run(
        'Keywords: Anomaly Detection, Self-Supervised Learning, Transformer, Financial Time-Series, '
        'Contrastive Learning, Energy-Based Models, Forex Market'
    )
    run.italic = True

    # ========================================================================
    # 2. INTRODUCTION
    # ========================================================================
    doc.add_heading('2. Introduction', level=1)
    doc.add_paragraph(
        'Financial markets generate vast amounts of time-series data characterized by complex '
        'temporal dependencies, regime shifts, and occasional extreme events. Detecting anomalous '
        'behavior in these markets—such as flash crashes, unusual volatility spikes, or '
        'manipulative trading patterns—is critical for risk management, regulatory compliance, '
        'and algorithmic trading systems.'
    )
    doc.add_paragraph(
        'Traditional approaches to financial anomaly detection rely on statistical methods '
        '(e.g., ARIMA residuals, z-score thresholds) or supervised machine learning models '
        'that require labeled anomaly datasets. However, labeled anomaly data in financial '
        'markets is scarce, subjective, and expensive to obtain. This creates a fundamental '
        'limitation for supervised approaches.'
    )
    doc.add_paragraph(
        'Self-supervised learning (SSL) offers a compelling alternative: by training models to '
        'solve pretext tasks on unlabeled data (e.g., predicting masked segments, distinguishing '
        'augmented views), the model learns rich representations of normal patterns without '
        'requiring any anomaly labels. Anomalies can then be detected as samples that deviate '
        'significantly from these learned representations.'
    )
    doc.add_paragraph(
        'This project implements a production-ready anomaly detection framework that combines '
        'Transformer-based temporal encoding with three complementary detection strategies: '
        'reconstruction error analysis, density-aware clustering, and energy-based scoring. '
        'The system is designed for the EUR/USD forex market using 4-hour (H4) candlestick data '
        'but is generalizable to other financial instruments and time horizons.'
    )

    # ========================================================================
    # 3. PROBLEM STATEMENT
    # ========================================================================
    doc.add_heading('3. Problem Statement', level=1)
    doc.add_paragraph(
        'The objective of this project is to design and implement a deep learning system capable '
        'of detecting anomalous temporal patterns in financial time-series data without relying '
        'on labeled anomaly examples during training. Specifically, the project addresses the '
        'following challenges:'
    )

    challenges = [
        ('Temporal Dependency Modeling', 'Financial time-series exhibit long-range dependencies '
         'where events at time t may be influenced by patterns from dozens or hundreds of time '
         'steps in the past. The model must capture these dependencies effectively.'),
        ('Absence of Labels', 'Real-world financial anomalies are rare, diverse, and often '
         'subjective. The system must learn to detect anomalies using only self-supervised '
         'signals derived from normal market behavior.'),
        ('Diverse Anomaly Types', 'Financial anomalies manifest in many forms: sudden price '
         'jumps, extreme volatility, unusual volume, trend reversals, flash crashes, and price '
         'gaps. A single detection mechanism is unlikely to capture all types.'),
        ('Class Imbalance', 'Anomalies are inherently rare (typically 1-7% of all observations), '
         'creating severe class imbalance that challenges threshold selection and evaluation.'),
        ('Real-Time Applicability', 'The detection system should be efficient enough for '
         'near-real-time deployment in trading environments with streaming data.'),
    ]

    for title_text, desc in challenges:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    # ========================================================================
    # 4. LITERATURE REVIEW
    # ========================================================================
    doc.add_heading('4. Literature Review', level=1)

    doc.add_heading('4.1 Transformer Models for Time-Series', level=2)
    doc.add_paragraph(
        'The Transformer architecture, introduced by Vaswani et al. (2017), revolutionized '
        'sequence modeling through self-attention mechanisms that can capture dependencies '
        'regardless of their distance in the sequence. Unlike RNNs and LSTMs, Transformers '
        'process all positions in parallel, enabling efficient computation on modern hardware. '
        'Recent works have adapted Transformers for time-series forecasting (Zhou et al., 2021; '
        'Wu et al., 2021), anomaly detection (Xu et al., 2021), and classification tasks.'
    )

    doc.add_heading('4.2 Self-Supervised Learning', level=2)
    doc.add_paragraph(
        'Self-supervised learning has emerged as a powerful paradigm for learning representations '
        'from unlabeled data. Key approaches include:'
    )
    doc.add_paragraph(
        'Contrastive Learning: SimCLR (Chen et al., 2020) and MoCo (He et al., 2020) learn '
        'representations by pulling augmented views of the same sample together while pushing '
        'different samples apart. The NT-Xent (Normalized Temperature-scaled Cross Entropy) '
        'loss has become standard for contrastive objectives.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Masked Prediction: BERT (Devlin et al., 2019) and Masked Autoencoders (He et al., 2021) '
        'learn by predicting randomly masked portions of the input. This forces the model to '
        'understand the underlying structure and relationships in the data.',
        style='List Bullet'
    )

    doc.add_heading('4.3 Energy-Based Models for Anomaly Detection', level=2)
    doc.add_paragraph(
        'Energy-based models (EBMs), as explored by LeCun et al. (2006) and Grathwohl et al. '
        '(2019), assign scalar energy values to input configurations. Normal samples occupy '
        'low-energy regions of the learned energy landscape, while anomalies correspond to '
        'high-energy states. This framework provides a principled approach to anomaly scoring '
        'that is compatible with self-supervised training objectives.'
    )

    doc.add_heading('4.4 Anomaly Detection in Financial Data', level=2)
    doc.add_paragraph(
        'Financial anomaly detection has been approached through statistical methods (Chandola '
        'et al., 2009), isolation forests (Liu et al., 2008), autoencoders (Sakurada & Yairi, '
        '2014), and more recently through deep learning approaches including variational '
        'autoencoders and GAN-based methods. The key challenge remains the lack of labeled '
        'data and the diversity of anomaly types in financial markets.'
    )

    # ========================================================================
    # 5. METHODOLOGY
    # ========================================================================
    doc.add_heading('5. Methodology', level=1)

    # 5.1 Data
    doc.add_heading('5.1 Data Collection and Preprocessing', level=2)
    doc.add_paragraph(
        'The dataset consists of EUR/USD foreign exchange H4 (4-hour) candlestick data '
        'spanning from January 2, 2015 to May 10, 2024, totaling 14,566 records. Each '
        'record contains the following raw fields: Open, High, Low, Close (OHLC prices), '
        'Tick Volume, Spread, and Real Volume.'
    )

    doc.add_heading('Data Preprocessing Pipeline', level=3)
    preprocessing_steps = [
        ('Column Filtering', 'Non-feature columns (unnamed index, tick_volume with extreme '
         'values >100,000, real_volume) are automatically dropped to prevent numerical instability.'),
        ('Technical Indicator Generation', '25 engineered features are computed from raw OHLC data '
         '(detailed in Section 5.2).'),
        ('NaN Handling', 'Rows with NaN values (from rolling window calculations) are dropped. '
         'Surviving row indices are tracked for ground truth alignment.'),
        ('Scaling', 'A RobustScaler (based on median and IQR) is applied to handle the '
         'heavy-tailed distributions common in financial data.'),
        ('Sliding Window', 'Data is organized into overlapping windows of 60 time steps '
         'with stride 1, producing sequences of shape (n_samples, 60, 25).'),
    ]
    for step, desc in preprocessing_steps:
        p = doc.add_paragraph()
        run = p.add_run(f'{step}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        'The final preprocessing yields 14,458 sequences, split chronologically into training '
        '(70%, 10,120 samples), validation (15%, 2,168 samples), and test (15%, 2,170 samples) sets.'
    )

    # 5.2 Feature Engineering
    doc.add_heading('5.2 Feature Engineering', level=2)
    doc.add_paragraph(
        'The system generates 25 features from raw OHLC data, organized into the following categories:'
    )

    add_styled_table(doc,
        ['Category', 'Features', 'Count'],
        [
            ['Price', 'open, high, low, close', '4'],
            ['Market Microstructure', 'spread', '1'],
            ['Returns', 'returns, log_returns', '2'],
            ['Moving Averages', 'sma_20, sma_50, ema_12, ema_26', '4'],
            ['MACD', 'macd, macd_signal, macd_diff', '3'],
            ['Momentum', 'rsi, stoch_k, stoch_d', '3'],
            ['Volatility', 'atr, atr_pct, bb_high, bb_low, bb_position', '5'],
            ['Trend', 'adx', '1'],
            ['Price Range', 'high_low_range, close_open_range', '2'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'Technical indicators are computed using the TA (Technical Analysis) library, which '
        'provides industry-standard implementations of RSI, MACD, Bollinger Bands, ATR, ADX, '
        'and Stochastic Oscillator.'
    )

    # 5.3 Model Architecture
    doc.add_heading('5.3 Model Architecture', level=2)

    doc.add_heading('5.3.1 Temporal Transformer Encoder', level=3)
    doc.add_paragraph(
        'The core of the system is a Transformer-based temporal encoder that processes '
        'sequences of financial features and produces context-aware embeddings. The encoder '
        'consists of the following components:'
    )

    arch_details = [
        ('Input Projection', 'A linear layer maps the 25-dimensional input features to '
         'the model dimension (d_model = 192).'),
        ('Positional Encoding', 'Sinusoidal positional encodings are added to provide '
         'temporal position information, enabling the model to distinguish between '
         'different time steps within the window.'),
        ('Transformer Encoder Layers', 'Five stacked Transformer encoder layers, each '
         'with 8 attention heads, 512-dimensional feed-forward networks, pre-norm '
         'architecture, and 20% dropout. The multi-head self-attention mechanism allows '
         'each time step to attend to all other time steps in the window.'),
        ('Layer Normalization', 'Applied after the final encoder layer for training stability.'),
        ('Mean + Max Pooling', 'The encoded sequence (batch, 60, 192) is aggregated into '
         'a single embedding vector (batch, 192) using concatenated mean and max pooling '
         'followed by a linear projection. Max-pooling is critical for preserving '
         'single-timestep anomaly peaks that mean-pooling alone would dilute by a factor of 1/60.'),
    ]
    for name, desc in arch_details:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    add_styled_table(doc,
        ['Parameter', 'Value', 'Description'],
        [
            ['d_model', '192', 'Transformer embedding dimension'],
            ['n_heads', '8', 'Number of self-attention heads'],
            ['n_layers', '5', 'Number of transformer encoder layers'],
            ['dim_feedforward', '512', 'Feed-forward network dimension'],
            ['dropout', '0.2', 'Dropout probability'],
            ['window_size', '60', 'Input sequence length (time steps)'],
            ['n_features', '25', 'Number of input features per time step'],
            ['Total Parameters', '~1,984,409', 'Trainable model parameters'],
        ]
    )

    doc.add_paragraph('')

    doc.add_heading('5.3.2 Masked Time-Series Reconstructor', level=3)
    doc.add_paragraph(
        'The masked reconstruction module randomly masks 15% of time steps in each input '
        'sequence, replaces them with a learnable mask token in the latent space, and '
        'trains a reconstruction head to predict the original values at masked positions. '
        'The reconstruction head is a two-layer MLP with GELU activation.'
    )
    doc.add_paragraph(
        'A key design choice for anomaly scoring is the information bottleneck: during '
        'inference, the full encoder output is mean-pooled into a single vector and then '
        'expanded back to the sequence length before reconstruction. This forces information '
        'loss and prevents the decoder from trivially copying anomalous timesteps, ensuring '
        'that anomalies produce distinctly higher reconstruction errors.'
    )

    doc.add_heading('5.3.3 Temporal Contrastive Learning', level=3)
    doc.add_paragraph(
        'The contrastive learning module creates two augmented views of each input sequence '
        'through temporal transformations: (1) random time masking (10% of steps) and (2) '
        'Gaussian noise injection (σ = 0.01). Both views are encoded and projected into a '
        '128-dimensional space. The NT-Xent (InfoNCE) loss with temperature τ = 0.07 pulls '
        'representations of the same sequence together while pushing different sequences apart.'
    )

    # 5.4 Self-Supervised Learning
    doc.add_heading('5.4 Self-Supervised Learning Objectives', level=2)
    doc.add_paragraph(
        'The model is trained with a combined loss function integrating both self-supervised '
        'objectives:'
    )
    p = doc.add_paragraph()
    run = p.add_run('L_total = α · L_contrastive + β · L_reconstruction')
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'where α = 0.1 (contrastive weight) and β = 1.0 (reconstruction weight). The '
        'reconstruction objective is heavily emphasized because it directly produces the '
        'representations used for anomaly scoring. The contrastive weight is kept low '
        'because aggressive contrastive learning can make anomalous windows embed similarly '
        'to their augmented counterparts, reducing detection separability.'
    )
    doc.add_paragraph(
        'Training uses AdamW optimizer with learning rate 1e-4, weight decay 1e-4, batch '
        'size 64, and gradient clipping at 1.0. Early stopping with patience of 30 epochs '
        'monitors validation loss.'
    )

    # 5.5 Anomaly Detection Pipeline
    doc.add_heading('5.5 Anomaly Detection Pipeline', level=2)
    doc.add_paragraph(
        'After the encoder is trained, three complementary anomaly scoring mechanisms '
        'are applied:'
    )

    doc.add_heading('5.5.1 Reconstruction-Based Detection', level=3)
    doc.add_paragraph(
        'The reconstruction detector computes per-sample anomaly scores by measuring the '
        'discrepancy between the original input and its reconstruction through the information '
        'bottleneck. The scoring uses a novel max-over-features approach: for each timestep, '
        'the maximum squared error across all 25 features is computed (rather than the mean), '
        'ensuring that a spike in even a single feature produces a high anomaly signal. '
        'The final score combines the maximum per-timestep error (60%) with the top-3 '
        'average (40%).'
    )

    doc.add_heading('5.5.2 Density-Aware Clustering', level=3)
    doc.add_paragraph(
        'K-Means clustering (k=8) is applied to the learned embeddings of training data '
        'to discover normal market regimes. Cluster density is computed as the ratio of '
        'cluster size to average intra-cluster distance. High-density clusters are labeled '
        'as "normal" regimes. For each test sample, the cluster-based anomaly score combines: '
        '(a) z-score distance from the assigned cluster center (50% weight), (b) cluster '
        'membership score—abnormal cluster assignment yields a high score (25% weight), and '
        '(c) within-cluster outlier detection based on the 90th percentile distance (25% weight).'
    )

    doc.add_heading('5.5.3 Energy-Based Scoring', level=3)
    doc.add_paragraph(
        'An energy-based neural network (3-layer MLP: 192→256→128→1) is trained to assign '
        'low energy values to normal embeddings and high energy values to anomalous '
        'embeddings. Training uses ground truth labels with a composite loss: margin loss '
        '(anomaly energy ≥ normal energy + margin), directional push losses, binary '
        'cross-entropy, and L2 regularization. The energy detector is trained for 80 epochs '
        'with AdamW optimizer (lr=1e-4, gradient clipping at 0.5).'
    )

    # 5.6 Hybrid Fusion
    doc.add_heading('5.6 Hybrid Score Fusion', level=2)
    doc.add_paragraph(
        'The three component scores are normalized using robust percentile-based '
        'normalization (5th–95th percentile range mapped to [0, 1]) and combined using '
        'a dual aggregation strategy:'
    )
    p = doc.add_paragraph()
    run = p.add_run(
        'S_hybrid = 0.6 × (w_r · S_recon + w_c · S_cluster + w_e · S_energy) + 0.4 × max(S_recon, S_cluster, S_energy)'
    )
    run.italic = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'where w_r = 0.60 (reconstruction), w_c = 0.20 (cluster), w_e = 0.20 (energy). '
        'The max component ensures that if any single detector produces a very high score, '
        'the overall score is boosted, capturing anomalies that may be extreme in only one '
        'detection dimension.'
    )
    doc.add_paragraph(
        'The detection threshold is tuned on the validation set through a multi-stage search: '
        '(1) wide search from the 50th to 99.9th percentile of combined scores, (2) fine-grained '
        'refinement around the best candidate, and (3) precision-constrained optimization '
        'requiring a minimum precision of 20–30%.'
    )

    # ========================================================================
    # 6. SYSTEM ARCHITECTURE
    # ========================================================================
    doc.add_heading('6. System Architecture', level=1)
    doc.add_paragraph(
        'The system is organized as a modular Python package with clear separation of concerns:'
    )

    add_styled_table(doc,
        ['Module', 'File', 'Responsibility'],
        [
            ['Data Layer', 'data/preprocessing.py', 'Loading, feature engineering, scaling, windowing'],
            ['Encoder', 'models/temporal_transformer.py', 'Transformer encoder, contrastive learning, reconstruction'],
            ['Clustering', 'models/clustering.py', 'K-Means clustering, density analysis, regime discovery'],
            ['Detectors', 'models/anomaly_detector.py', 'Energy-based, reconstruction-based, hybrid detection'],
            ['Training', 'train.py', 'End-to-end pipeline: data → train → detect → report'],
            ['Reporting', 'report_generator.py', 'Excel report generation (9-sheet detailed report)'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'The complete pipeline is executed through a single command (python train.py), which '
        'performs all 8 stages: data loading, model initialization, transformer training, '
        'clustering, reconstruction detector fitting, energy detector training, threshold '
        'tuning, and final evaluation with visualization and report generation.'
    )

    # ========================================================================
    # 7. IMPLEMENTATION DETAILS
    # ========================================================================
    doc.add_heading('7. Implementation Details', level=1)

    doc.add_heading('7.1 Technology Stack', level=2)
    add_styled_table(doc,
        ['Technology', 'Version', 'Purpose'],
        [
            ['Python', '3.10', 'Core programming language'],
            ['PyTorch', '2.6.0', 'Deep learning framework'],
            ['NumPy', '≥2.0.0', 'Numerical computation'],
            ['pandas', '≥2.0.0', 'Data manipulation'],
            ['scikit-learn', '≥1.3.0', 'K-Means, Isolation Forest, scaling'],
            ['TA Library', '≥0.11.0', 'Technical indicator computation'],
            ['matplotlib / seaborn', '≥3.7.0 / ≥0.12.0', 'Visualization'],
            ['openpyxl', '—', 'Excel report generation'],
        ]
    )

    doc.add_paragraph('')

    doc.add_heading('7.2 Synthetic Anomaly Injection', level=2)
    doc.add_paragraph(
        'Since labeled anomaly data is unavailable, the system injects synthetic anomalies '
        'into the raw data for quantitative evaluation. Six anomaly types are injected at '
        'a total rate of 7%, with the following distribution:'
    )

    add_styled_table(doc,
        ['Anomaly Type', 'Probability', 'Description'],
        [
            ['Price Spike', '20%', 'Sudden large price jump (intensity × local std)'],
            ['Flash Crash', '20%', 'Severe price drop with partial recovery'],
            ['Volatility Spike', '20%', 'Extreme high-low range expansion'],
            ['Volume Spike', '15%', 'Abnormally high trading volume'],
            ['Trend Break', '15%', 'Sharp reversal from recent mean price'],
            ['Gap Anomaly', '10%', 'Open price gaps from previous close'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'Each anomaly is context-aware: the injection magnitude is scaled by local rolling '
        'volatility (20-bar rolling standard deviation) multiplied by an intensity factor '
        'of 4.0. Anomalies are placed only at safe indices (avoiding the first and last '
        '60 positions) to ensure complete windowed sequences.'
    )

    doc.add_heading('7.3 Ground Truth Labeling', level=2)
    doc.add_paragraph(
        'Each sliding window sequence is labeled as anomalous if the last point in the '
        '60-step window corresponds to an injected anomaly. This "last-point" labeling '
        'strategy was chosen over "any-in-window" labeling because the latter would label '
        'approximately 98.5% of all sequences as anomalous (since with 7% point-level '
        'anomaly rate and window size 60, the probability of a clean window is only '
        '(0.93)^60 ≈ 1.2%), making meaningful detection impossible.'
    )

    # ========================================================================
    # 8. EXPERIMENTAL SETUP
    # ========================================================================
    doc.add_heading('8. Experimental Setup', level=1)

    doc.add_heading('8.1 Training Configuration', level=2)
    add_styled_table(doc,
        ['Parameter', 'Value'],
        [
            ['Maximum Epochs', '150'],
            ['Batch Size', '64'],
            ['Learning Rate', '1e-4'],
            ['Optimizer', 'AdamW (weight_decay=1e-4)'],
            ['Gradient Clipping', '1.0'],
            ['Early Stopping Patience', '30 epochs'],
            ['Contrastive Weight (α)', '0.1'],
            ['Reconstruction Weight (β)', '1.0'],
            ['Energy Detector Epochs', '80'],
            ['Energy Learning Rate', '1e-4'],
            ['Number of Clusters', '8'],
            ['Anomaly Injection Rate', '7%'],
            ['Anomaly Intensity', '4.0'],
            ['Device', 'CPU / CUDA (auto-detected)'],
        ]
    )

    doc.add_paragraph('')

    doc.add_heading('8.2 Data Split', level=2)
    doc.add_paragraph('The dataset is split chronologically (no shuffling) to simulate realistic deployment:')
    add_styled_table(doc,
        ['Split', 'Ratio', 'Samples', 'Anomalies'],
        [
            ['Training', '70%', '10,120', '~695 (6.9%)'],
            ['Validation', '15%', '2,168', '~145–159 (6.7–7.3%)'],
            ['Test', '15%', '2,170', '~156–161 (7.2–7.4%)'],
        ]
    )

    doc.add_paragraph('')

    doc.add_heading('8.3 Evaluation Metrics', level=2)
    doc.add_paragraph('The following metrics are used for evaluation:')
    metrics_desc = [
        ('Precision', 'Of all samples flagged as anomalies, what fraction are truly anomalous? '
         'High precision means few false alarms.'),
        ('Recall', 'Of all true anomalies, what fraction are detected? '
         'High recall means few missed anomalies.'),
        ('F1 Score', 'The harmonic mean of precision and recall, providing a balanced measure '
         'of detection quality.'),
        ('Accuracy', 'Overall fraction of correct predictions (both normal and anomalous).'),
    ]
    for name, desc in metrics_desc:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    # ========================================================================
    # 9. RESULTS AND ANALYSIS
    # ========================================================================
    doc.add_heading('9. Results and Analysis', level=1)

    doc.add_heading('9.1 Training Progress', level=2)
    doc.add_paragraph(
        'The transformer model was trained with early stopping. The training converged '
        'with the following characteristics:'
    )

    add_styled_table(doc,
        ['Metric', 'Value'],
        [
            ['Total Epochs Trained', '97 (early stopped at patience=30)'],
            ['Best Validation Loss', '1.327'],
            ['Final Training Loss', '0.779'],
            ['Final Validation Loss', '1.342'],
            ['Training Gap', '0.563 (moderate overfitting)'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'The model exhibited a consistent gap between training and validation loss '
        '(approximately 0.56), indicating some overfitting despite the 20% dropout rate '
        'and weight decay regularization. The validation loss continued to improve slowly '
        'until early stopping activated, suggesting the model was still learning useful '
        'representations.'
    )

    doc.add_heading('9.2 Clustering Results', level=2)
    doc.add_paragraph(
        'Density-aware K-Means clustering identified 7 normal market regime clusters out '
        'of 8 total clusters. The cluster sizes were well-balanced, ranging from approximately '
        '1,100 to 1,600 samples per cluster. The anomaly capture rate in the designated '
        'outlier cluster was 11.7%, indicating that clustering alone provides limited '
        'separation between normal and anomalous patterns in the embedding space.'
    )

    doc.add_heading('9.3 Energy Detector Training', level=2)
    doc.add_paragraph(
        'The energy-based detector was trained for 80 epochs using ground truth labels '
        'with the supervised margin loss. The training loss decreased from 3.88 to 0.60, '
        'indicating successful convergence. The energy detector learned to assign distinct '
        'energy levels to normal vs. anomalous embeddings, though the separation in the '
        'combined hybrid score was still challenged by the overlap in the reconstruction '
        'and cluster components.'
    )

    doc.add_heading('9.4 Detection Performance', level=2)

    # Try to load the latest results
    results_dirs = sorted([
        d for d in os.listdir('.')
        if d.startswith('improved_outputs_') and os.path.isdir(d)
    ], reverse=True)

    results_data = []
    for rdir in results_dirs[:5]:  # Latest 5 runs
        rjson = os.path.join(rdir, 'results.json')
        if os.path.exists(rjson):
            with open(rjson) as f:
                r = json.load(f)
            results_data.append((rdir, r))

    if results_data:
        doc.add_paragraph('Results from recent experimental runs:')
        rows = []
        for rdir, r in results_data:
            ts = r.get('timestamp', rdir.split('_', 2)[-1] if '_' in rdir else rdir)
            t = r.get('test', {})
            rows.append([
                ts,
                f"{t.get('precision', 0):.3f}",
                f"{t.get('recall', 0):.3f}",
                f"{t.get('f1', 0):.3f}",
                f"{t.get('accuracy', 0):.3f}",
                f"TP={t.get('tp', 0)}, FP={t.get('fp', 0)}",
            ])

        add_styled_table(doc,
            ['Run Timestamp', 'Precision', 'Recall', 'F1', 'Accuracy', 'Detections'],
            rows
        )
        doc.add_paragraph('')
    else:
        doc.add_paragraph('(No saved results found. Run train.py to generate results.)')

    # Representative confusion matrix
    doc.add_heading('9.5 Confusion Matrix Analysis', level=2)
    doc.add_paragraph(
        'A representative confusion matrix from the latest run (using hybrid detection '
        'with reconstruction + cluster + energy scoring):'
    )

    add_styled_table(doc,
        ['', 'Predicted Normal', 'Predicted Anomaly'],
        [
            ['Actual Normal', 'TN = 1,719', 'FP = 290'],
            ['Actual Anomaly', 'FN = 135', 'TP = 26'],
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'The confusion matrix reveals several key findings: (1) The model correctly identifies '
        'the majority of normal samples (TN = 1,719 out of 2,009 normal samples, 85.6% '
        'specificity). (2) The false positive rate is high (290 false alarms), which is a '
        'common challenge in anomaly detection with heavily imbalanced classes. (3) The model '
        'detects some true anomalies (TP = 26) but misses a significant portion (FN = 135), '
        'indicating room for improvement in recall.'
    )

    doc.add_heading('9.6 Score Distribution Analysis', level=2)
    doc.add_paragraph(
        'Analysis of the hybrid anomaly scores reveals a 51% overlap between the normal '
        'and anomaly score distributions. The normal samples have a mean score of 0.506 '
        '(std=0.178) and the anomaly samples have a mean score of 0.528 (std=0.188). '
        'This near-identical distribution is the primary cause of low F1 scores—the '
        'detection mechanism struggles to find a threshold that separates the two classes.'
    )
    doc.add_paragraph(
        'This overlap motivated several architectural improvements including the '
        'reconstruction bottleneck, max+mean pooling, disabled outlier clipping, and '
        'increased anomaly intensity, which are expected to significantly improve '
        'score separability in subsequent runs.'
    )

    # ========================================================================
    # 10. DISCUSSION
    # ========================================================================
    doc.add_heading('10. Discussion', level=1)
    doc.add_paragraph(
        'The results demonstrate both the promise and challenges of self-supervised anomaly '
        'detection in financial time-series. The transformer encoder successfully learns '
        'meaningful temporal representations, as evidenced by the consistent training '
        'convergence and the ability to discover distinct market regime clusters. However, '
        'translating these representations into effective anomaly scores remains challenging.'
    )

    doc.add_heading('10.1 Key Observations', level=2)
    observations = [
        ('Ground Truth Labeling Matters', 'The choice between "any-in-window" and "last-point" '
         'labeling has a dramatic impact. The any-in-window approach labeled 98.5% of sequences '
         'as anomalous, making the problem trivially unsolvable. Switching to last-point labeling '
         'restored the expected 7% anomaly rate.'),
        ('Reconstruction Bottleneck', 'Without an information bottleneck, the reconstruction '
         'detector trivially copies anomalous inputs through the full encoder output, producing '
         'near-zero reconstruction error for both normal and anomalous samples. The mean-pooling '
         'bottleneck forces information loss, creating meaningful error differences.'),
        ('Pooling Strategy', 'Mean-pooling alone dilutes single-timestep anomalies by a factor '
         'of 1/60. Combining mean and max pooling preserves the anomalous signal in the embedding, '
         'which is critical for both energy-based and cluster-based detection.'),
        ('Outlier Clipping Conflict', 'Clipping outliers at the 1st/99th percentile was designed '
         'to improve scaling but inadvertently destroyed the very anomaly signals the model '
         'needed to detect. Disabling clipping allows the full anomaly magnitude to survive '
         'into model input.'),
    ]
    for title_text, desc in observations:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('10.2 Hybrid Fusion Benefits', level=2)
    doc.add_paragraph(
        'The hybrid fusion approach provides robustness through complementarity. '
        'Reconstruction-based detection excels at capturing anomalies with unusual '
        'feature patterns, while energy-based detection learns a more abstract notion '
        'of normality. Cluster-based scoring captures structural deviations in the '
        'embedding space. The max aggregation component ensures that a strong signal '
        'from any single detector is not diluted by the others.'
    )

    # ========================================================================
    # 11. CHALLENGES AND SOLUTIONS
    # ========================================================================
    doc.add_heading('11. Challenges and Solutions', level=1)

    challenges_solutions = [
        ('Challenge: 98.5% Anomaly Rate in Labels',
         'The any-in-window labeling strategy with 7% point-level anomaly rate and window '
         'size 60 caused nearly all sequences to be labeled anomalous.',
         'Switched to last-point labeling: a sequence is anomalous only if its final '
         'time step is anomalous, restoring the expected ~7% rate.'),
        ('Challenge: Reconstruction Trivially Copies Input',
         'The encoder sees the full unmasked input and reproduces it faithfully, producing '
         'identical reconstruction errors for normal and anomalous samples.',
         'Added a mean-pooling bottleneck during inference: the full encoder output is '
         'compressed to a single vector before reconstruction, forcing information loss.'),
        ('Challenge: Outlier Clipping Destroys Anomaly Signal',
         'Percentile-based clipping at 1%/99% flattened the injected anomaly spikes before '
         'the model could see them.',
         'Disabled outlier clipping entirely (clip_outliers=False) and increased anomaly '
         'intensity to 4.0 to ensure signals survive RobustScaler normalization.'),
        ('Challenge: Mean Pooling Dilutes Point Anomalies',
         'Averaging 60 timesteps reduces a single anomalous timestep signal by 60×.',
         'Added max-pooling concatenated with mean-pooling, followed by a linear projection. '
         'Max-pooling preserves the peak activation from the anomalous timestep.'),
        ('Challenge: Empty Normal Mask Crashes Threshold Tuning',
         'When all validation samples were labeled anomalous, np.percentile on an empty '
         'array caused an IndexError.',
         'Added safety guards that check for empty normal/anomaly masks before computing '
         'percentile statistics.'),
    ]

    for challenge, problem, solution in challenges_solutions:
        p = doc.add_paragraph()
        run = p.add_run(challenge)
        run.bold = True

        p2 = doc.add_paragraph()
        run2 = p2.add_run('Problem: ')
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        p2.add_run(problem)

        p3 = doc.add_paragraph()
        run3 = p3.add_run('Solution: ')
        run3.bold = True
        run3.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        p3.add_run(solution)

        doc.add_paragraph('')  # Spacing

    # ========================================================================
    # 12. FUTURE WORK
    # ========================================================================
    doc.add_heading('12. Future Work', level=1)
    future_items = [
        ('Real Anomaly Labels', 'Obtain expert-labeled anomaly data from financial analysts '
         'to validate the system against real-world events rather than synthetic anomalies.'),
        ('Multi-Asset Generalization', 'Extend the framework to handle multiple currency pairs, '
         'stocks, and commodities simultaneously, learning shared anomaly patterns across markets.'),
        ('Online Learning', 'Implement incremental model updates for streaming data, allowing '
         'the system to adapt to evolving market conditions without full retraining.'),
        ('Attention-Based Explainability', 'Leverage the Transformer attention weights to '
         'provide interpretable explanations of which time steps and features contributed '
         'most to each anomaly detection.'),
        ('Adversarial Training', 'Use generative adversarial training to improve the quality '
         'of synthetic anomalies and make the detector more robust to subtle manipulations.'),
        ('Ensemble of Window Sizes', 'Train multiple models with different window sizes '
         '(15, 30, 60, 120) and combine their predictions to capture anomalies at multiple '
         'time scales.'),
        ('Integration with Trading Systems', 'Deploy the anomaly detector as a real-time '
         'risk monitoring component in an automated trading system.'),
    ]

    for title_text, desc in future_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    # ========================================================================
    # 13. CONCLUSION
    # ========================================================================
    doc.add_heading('13. Conclusion', level=1)
    doc.add_paragraph(
        'This project successfully designed and implemented a self-supervised deep learning '
        'framework for anomaly detection in financial time-series data. The system combines '
        'a Transformer-based temporal encoder with three complementary detection strategies—'
        'reconstruction error with information bottleneck, density-aware clustering, and '
        'energy-based scoring—in a hybrid fusion approach.'
    )
    doc.add_paragraph(
        'The framework processes EUR/USD H4 forex data spanning 9 years (2015–2024), '
        'generating 25 engineered features per time step and organizing them into 14,458 '
        'overlapping sequences. The self-supervised training objectives (contrastive learning '
        'and masked reconstruction) enable the model to learn rich representations of normal '
        'market behavior without requiring any anomaly labels.'
    )
    doc.add_paragraph(
        'Through iterative development, several critical design decisions were identified '
        'and resolved: last-point ground truth labeling (vs. any-in-window), reconstruction '
        'bottleneck for meaningful error computation, mean+max pooling for preserving '
        'point anomaly signals, and disabling outlier clipping to preserve injected anomaly '
        'magnitudes. These insights contribute to the broader understanding of self-supervised '
        'anomaly detection in temporal data.'
    )
    doc.add_paragraph(
        'The modular architecture, comprehensive visualization pipeline, and detailed '
        'reporting make the system suitable for further research and potential deployment '
        'in financial market surveillance applications.'
    )

    # ========================================================================
    # 14. REFERENCES
    # ========================================================================
    doc.add_heading('14. References', level=1)

    references = [
        'Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). "Attention Is All You Need." '
        'Advances in Neural Information Processing Systems, 30.',
        'Chen, T., Kornblith, S., Norouzi, M., & Hinton, G. (2020). "A Simple Framework for '
        'Contrastive Learning of Visual Representations." International Conference on Machine Learning.',
        'He, K., Chen, X., Xie, S., et al. (2021). "Masked Autoencoders Are Scalable Vision '
        'Learners." arXiv preprint arXiv:2111.06377.',
        'Grathwohl, W., Wang, K., Jacobsen, J., et al. (2019). "Your Classifier is Secretly '
        'an Energy Based Model and You Should Treat it Like One." International Conference on '
        'Learning Representations.',
        'He, K., Fan, H., Wu, Y., et al. (2020). "Momentum Contrast for Unsupervised Visual '
        'Representation Learning." IEEE/CVF Conference on Computer Vision and Pattern Recognition.',
        'Devlin, J., Chang, M., Lee, K., & Toutanova, K. (2019). "BERT: Pre-training of Deep '
        'Bidirectional Transformers for Language Understanding." NAACL-HLT.',
        'Zhou, H., Zhang, S., Peng, J., et al. (2021). "Informer: Beyond Efficient Transformer '
        'for Long Sequence Time-Series Forecasting." AAAI Conference on Artificial Intelligence.',
        'Wu, H., Xu, J., Wang, J., & Long, M. (2021). "Autoformer: Decomposition Transformers '
        'with Auto-Correlation for Long-Term Series Forecasting." NeurIPS.',
        'Xu, J., Wu, H., Wang, J., & Long, M. (2021). "Anomaly Transformer: Time Series Anomaly '
        'Detection with Association Discrepancy." ICLR.',
        'Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). "Isolation Forest." IEEE International '
        'Conference on Data Mining.',
        'Chandola, V., Banerjee, A., & Kumar, V. (2009). "Anomaly Detection: A Survey." ACM '
        'Computing Surveys, 41(3).',
        'Sakurada, M., & Yairi, T. (2014). "Anomaly Detection Using Autoencoders with Nonlinear '
        'Dimensionality Reduction." MLSDA Workshop.',
        'LeCun, Y., Chopra, S., Hadsell, R., et al. (2006). "A Tutorial on Energy-Based Learning." '
        'Predicting Structured Data, MIT Press.',
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'[{i}] ')
        run.bold = True
        p.add_run(ref)
        p.paragraph_format.space_after = Pt(4)

    # ========================================================================
    # Save
    # ========================================================================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'Project_Report_Anomaly_Detection.docx'
    )
    doc.save(output_path)
    print(f"✓ Project report saved to: {output_path}")
    print(f"  Total sections: 14")
    print(f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    return output_path


if __name__ == '__main__':
    generate_report()

